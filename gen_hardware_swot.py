"""
Generate hardware_swot.docx — SWOT analysis for each hardware tier
in the Geo-Agentic RAG local pipeline recommendation.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
C_DARK    = RGBColor(0x1E, 0x29, 0x3B)   # dark navy  (header bg)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)   # white text
C_GREY_BG = RGBColor(0xF1, 0xF5, 0xF9)  # light slate (alternate row)

# SWOT quadrant colours
S_BG = RGBColor(0x16, 0xA3, 0x4A)   # green  — Strengths
W_BG = RGBColor(0xDC, 0x26, 0x26)   # red    — Weaknesses
O_BG = RGBColor(0x25, 0x63, 0xEB)   # blue   — Opportunities
T_BG = RGBColor(0xD9, 0x77, 0x06)   # amber  — Threats

# ── Hardware tier data ────────────────────────────────────────────────────────
TIERS = [
    {
        "tier": "Tier 1 — Tight but Works (Current Setup)",
        "subtitle": "~4 GB VRAM GPU",
        "specs": [
            ("GPU",          "Any 4 GB VRAM card (e.g., GTX 1650, RTX 3050)"),
            ("VRAM",         "4 GB"),
            ("System RAM",   "8–16 GB DDR4"),
            ("Storage",      "256 GB SSD (model files)"),
            ("CPU",          "Any modern quad-core"),
            ("LLM Model",    "Qwen2.5:3b (4-bit quant, ~2.0 GB VRAM)"),
            ("Concurrency",  "OLLAMA_NUM_PARALLEL=2 → 3–5 users"),
        ],
        "swot": {
            "S": [
                "Zero additional cost — uses existing hardware",
                "Fits entire RAG stack (3b + e5-large + bge-reranker ~3.6 GB)",
                "Suitable for solo development and unit testing",
                "Fast iteration — no cloud latency",
            ],
            "W": [
                "Qwen2.5:3b produces noticeably weaker answers than 7b/14b",
                "Only 3–5 concurrent users before queuing",
                "Insufficient for live 120-respondent SUS evaluation",
                "No VRAM headroom for model upgrades",
            ],
            "O": [
                "Enough for all development milestones and RAG tuning",
                "Can still measure Recall@K / MRR / Precision@K for dissertation metrics",
                "Low barrier — start building today without any purchase",
            ],
            "T": [
                "VRAM OOM if all three models loaded simultaneously at peak load",
                "Bottleneck exposed if used during actual respondent evaluation",
                "Dissertation findings may reflect model quality limit, not system design",
            ],
        },
    },
    {
        "tier": "Tier 2 — Dissertation-Grade (Recommended Upgrade)",
        "subtitle": "RTX 3060 12 GB  /  RTX 4060 Ti 16 GB",
        "specs": [
            ("GPU",          "RTX 3060 12 GB (~₱15k–20k used)"),
            ("VRAM",         "12 GB"),
            ("System RAM",   "16 GB DDR4"),
            ("Storage",      "256–512 GB NVMe"),
            ("CPU",          "Ryzen 5 / Intel i5, 6-core"),
            ("LLM Model",    "Qwen2.5:7b (4-bit quant, ~4.5 GB VRAM)"),
            ("Concurrency",  "OLLAMA_NUM_PARALLEL=3 → 10–15 users"),
        ],
        "swot": {
            "S": [
                "Fits Qwen2.5:7b + e5-large + bge-reranker (~6.1 GB, 5.9 GB headroom)",
                "7b model measurably improves Recall@K — directly strengthens dissertation results",
                "Handles 120-respondent SUS evaluation entirely offline",
                "Best price-to-performance ratio for research use",
                "No cloud dependency, no rate limits, no per-query cost",
            ],
            "W": [
                "Upfront cost (~₱15k–20k for used RTX 3060)",
                "Still limited to ~10–15 concurrent users under heavy load",
                "vLLM requires WSL2 setup for PagedAttention throughput gains",
            ],
            "O": [
                "7b quality gap vs 3b is publishable — compare both in dissertation evaluation",
                "Upgrade path clear: swap to 14b model if 12 GB VRAM allows (~10.6 GB with 14b)",
                "Doubles as general ML research workstation for future projects",
                "Can demo to city hall stakeholders without internet",
            ],
            "T": [
                "Used GPU market: verify VRAM integrity before purchase",
                "Power consumption increases (~170 W TDP vs ~75 W for entry cards)",
                "Driver/CUDA updates may require maintenance overhead",
            ],
        },
    },
    {
        "tier": "Tier 3 — Deployment-Grade",
        "subtitle": "RTX 3090 24 GB  /  RTX 4090 24 GB",
        "specs": [
            ("GPU",          "RTX 3090 24 GB (~₱30k–40k used)"),
            ("VRAM",         "24 GB"),
            ("System RAM",   "32 GB DDR4/DDR5"),
            ("Storage",      "512 GB NVMe"),
            ("CPU",          "Ryzen 7 / Intel i7, 8-core"),
            ("LLM Model",    "Qwen2.5:14b (4-bit quant, ~9.0 GB VRAM)"),
            ("Concurrency",  "OLLAMA_NUM_PARALLEL=4–6 → 20–40 users; vLLM → 40–60 users"),
        ],
        "swot": {
            "S": [
                "Qwen2.5:14b fits with 13 GB VRAM to spare — highest local answer quality",
                "40–60 concurrent users with vLLM PagedAttention on WSL2",
                "Viable for actual city hall department deployment (1–2 counters)",
                "24 GB future-proofs against larger models (32b at ~18 GB 4-bit)",
                "Full analytics pipeline runs unthrottled",
            ],
            "W": [
                "High upfront cost (~₱30k–40k used)",
                "RTX 3090 has 350 W TDP — significant electricity cost at 24/7 operation",
                "vLLM requires WSL2 + Linux environment; adds maintenance complexity",
                "Overkill for dissertation evaluation alone",
            ],
            "O": [
                "City hall pilot deployment without any cloud subscription",
                "Can run Qwen2.5:32b at 4-bit (~18 GB) — near GPT-3.5 quality locally",
                "Enables parallel multi-model A/B testing for RAG quality studies",
                "Strong platform for future NLP research projects",
            ],
            "T": [
                "RTX 3090 VRAM uses GDDR6X — known to run hot; needs adequate airflow",
                "Power bill at 350 W × 24 h: ~₱2,500–3,000/month at Philippine rates",
                "Single point of failure — no redundancy for production government use",
            ],
        },
    },
    {
        "tier": "Tier 4 — No-Compromise Production",
        "subtitle": "Mac Mini M4 Pro  /  Mac Studio M3 Ultra",
        "specs": [
            ("Device",        "Mac Mini M4 Pro (24 GB) or Mac Studio M3 Ultra (192 GB)"),
            ("Unified Memory","24 GB – 192 GB (shared CPU + GPU)"),
            ("System RAM",    "Same as unified memory — no separate allocation"),
            ("Storage",       "256 GB – 8 TB SSD (Apple internal)"),
            ("CPU",           "M4 Pro: 14-core / M3 Ultra: 32-core"),
            ("LLM Model",     "Qwen2.5:14b–72b depending on memory tier"),
            ("Concurrency",   "Mac Mini M4 Pro: ~30–50 users; M3 Ultra: ~200+ users"),
        ],
        "swot": {
            "S": [
                "Unified memory: entire model (LLM + encoders) in one flat pool — no VRAM limit",
                "Ollama runs natively on Apple Silicon — no WSL2, no Linux needed",
                "Silent, fanless-class operation — suitable for office deployment",
                "Low power: M4 Pro Mac Mini ~30 W idle, ~60 W load vs 350 W for 3090",
                "macOS stability — minimal driver/OS maintenance overhead",
                "Mac Mini M4 Pro at ₱60k–80k competitive with RTX 3090 builds for this workload",
            ],
            "W": [
                "High upfront cost, especially Mac Studio (₱150k–300k+)",
                "Non-upgradeable memory — must spec correctly at purchase",
                "GPU compute throughput lower than equivalent NVIDIA VRAM tier for CUDA-only workloads",
                "No CUDA — some Python ML libraries have Apple Silicon workarounds, not native CUDA",
            ],
            "O": [
                "M3 Ultra 192 GB can run Qwen2.5:72b locally — GPT-4 class answers, zero cloud cost",
                "Perfect always-on city hall server: low power, no noise, high reliability",
                "Future Apple Silicon generations will only increase performance per watt",
                "macOS Server features simplify HTTPS certificate management",
            ],
            "T": [
                "Locked ecosystem — hardware failures require Apple authorised repair",
                "CUDA-dependent libraries (some FAISS optimisations) need Metal/CPU fallback",
                "City hall IT procurement may prefer Windows-based servers (policy risk)",
                "Resale market thinner than PC parts if hardware needs change",
            ],
        },
    },
]

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_col = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_col)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top/bottom/left/right = (size_pt, color_hex)"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, (sz, col) in kwargs.items():
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    str(sz * 8))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), col)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def para(cell, text, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet_para(cell, text, color=None, size=9.5):
    p   = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(f"• {text}")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def build_swot_table(doc, swot):
    """Build a 2×2 SWOT table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    quadrants = [
        (0, 0, "S  —  STRENGTHS",    swot["S"], S_BG),
        (0, 1, "W  —  WEAKNESSES",   swot["W"], W_BG),
        (1, 0, "O  —  OPPORTUNITIES",swot["O"], O_BG),
        (1, 1, "T  —  THREATS",      swot["T"], T_BG),
    ]

    col_width = Inches(3.15)
    for row_cells in tbl.rows:
        for cell in row_cells.cells:
            cell.width = col_width

    for r, c, label, items, bg in quadrants:
        cell = tbl.cell(r, c)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Header paragraph
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WHITE

        set_cell_bg(cell, bg)

        for item in items:
            add_bullet_para(cell, item, color=C_WHITE, size=9)

    # Remove outer border
    tbl.style = 'Table Grid'
    return tbl


# ── Main document build ───────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Title page area ──────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Hardware Specification SWOT Analysis")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = C_DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "Geo-Agentic RAG — Full Local Pipeline · No Rate Limits · No Cloud Cost"
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    sr.italic = True

    doc.add_paragraph()  # spacer

    # VRAM budget quick-ref table
    vram_hdr = doc.add_paragraph()
    vr = vram_hdr.add_run("VRAM Budget Reference")
    vr.bold = True
    vr.font.size = Pt(11)
    vr.font.color.rgb = C_DARK

    vram_rows = [
        ("Component",              "VRAM (4-bit quant)", True),
        ("multilingual-e5-large",  "~1.1 GB",           False),
        ("bge-reranker-base",      "~0.5 GB",           False),
        ("Qwen2.5:3b",             "~2.0 GB",           False),
        ("Qwen2.5:7b",             "~4.5 GB",           False),
        ("Qwen2.5:14b",            "~9.0 GB",           False),
        ("Qwen2.5:32b",            "~18.0 GB",          False),
        ("Full stack (3b model)",  "~3.6 GB total",     False),
        ("Full stack (7b model)",  "~6.1 GB total",     False),
        ("Full stack (14b model)", "~10.6 GB total",    False),
    ]

    vtbl = doc.add_table(rows=len(vram_rows), cols=2)
    vtbl.style = 'Table Grid'
    vtbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (comp, vram, is_hdr) in enumerate(vram_rows):
        row = vtbl.rows[i]
        for j, txt in enumerate([comp, vram]):
            cell = row.cells[j]
            cell.width = Inches(2.8) if j == 0 else Inches(1.8)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9)
            run.bold = is_hdr
            if is_hdr:
                set_cell_bg(cell, C_DARK)
                run.font.color.rgb = C_WHITE
            elif i % 2 == 0:
                set_cell_bg(cell, C_GREY_BG)

    doc.add_paragraph()  # spacer

    # ── One page per tier ────────────────────────────────────────────────────
    for idx, tier in enumerate(TIERS):
        if idx > 0:
            doc.add_page_break()

        # Tier heading
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = h.add_run(tier["tier"])
        hr.bold = True
        hr.font.size = Pt(14)
        hr.font.color.rgb = C_DARK

        sh = doc.add_paragraph()
        sr2 = sh.add_run(tier["subtitle"])
        sr2.italic = True
        sr2.font.size = Pt(10)
        sr2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()

        # Spec table
        spec_hdr = doc.add_paragraph()
        shr = spec_hdr.add_run("Hardware Specifications")
        shr.bold = True
        shr.font.size = Pt(10)
        shr.font.color.rgb = C_DARK

        stbl = doc.add_table(rows=len(tier["specs"]) + 1, cols=2)
        stbl.style = 'Table Grid'
        stbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        hrow = stbl.rows[0]
        for j, txt in enumerate(["Specification", "Detail"]):
            cell = hrow.cells[j]
            cell.width = Inches(1.8) if j == 0 else Inches(4.5)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = C_WHITE
            set_cell_bg(cell, C_DARK)

        for i, (spec, detail) in enumerate(tier["specs"]):
            row = stbl.rows[i + 1]
            for j, txt in enumerate([spec, detail]):
                cell = row.cells[j]
                cell.width = Inches(1.8) if j == 0 else Inches(4.5)
                p = cell.paragraphs[0]
                run = p.add_run(txt)
                run.font.size = Pt(9)
                run.bold = (j == 0)
                if i % 2 == 0:
                    set_cell_bg(cell, C_GREY_BG)

        doc.add_paragraph()

        # SWOT label
        swot_hdr = doc.add_paragraph()
        swhr = swot_hdr.add_run("SWOT Analysis")
        swhr.bold = True
        swhr.font.size = Pt(10)
        swhr.font.color.rgb = C_DARK

        build_swot_table(doc, tier["swot"])

    # ── Summary recommendation table ─────────────────────────────────────────
    doc.add_page_break()

    rec_title = doc.add_paragraph()
    rec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rtr = rec_title.add_run("Recommendation Summary")
    rtr.bold = True
    rtr.font.size = Pt(14)
    rtr.font.color.rgb = C_DARK

    doc.add_paragraph()

    rec_rows = [
        ("Goal / Scenario",                           "Recommended Tier",            True),
        ("Finish dissertation, solo development",     "Tier 1 — Current Setup",      False),
        ("120-respondent SUS evaluation (local)",     "Tier 2 — RTX 3060 12 GB",     False),
        ("City hall department pilot deployment",     "Tier 3 — RTX 3090 24 GB",     False),
        ("Always-on, low-power office server",        "Tier 4 — Mac Mini M4 Pro",    False),
        ("City-wide scale, max answer quality",       "Tier 4 — Mac Studio M3 Ultra",False),
    ]

    rtbl = doc.add_table(rows=len(rec_rows), cols=2)
    rtbl.style = 'Table Grid'
    rtbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (goal, rec, is_hdr) in enumerate(rec_rows):
        row = rtbl.rows[i]
        for j, txt in enumerate([goal, rec]):
            cell = row.cells[j]
            cell.width = Inches(3.2)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9.5)
            run.bold = is_hdr
            if is_hdr:
                set_cell_bg(cell, C_DARK)
                run.font.color.rgb = C_WHITE
            elif i % 2 == 0:
                set_cell_bg(cell, C_GREY_BG)

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        "Note: All Philippine Peso estimates (₱) are approximate as of 2025 and reflect "
        "used/second-hand market prices. VRAM figures assume 4-bit quantisation via Ollama's "
        "default GGUF backend. Concurrency figures assume the full 6-stage RAG pipeline "
        "(query rewrite → multilingual-e5-large → FAISS → bge-reranker → LLM → coordinate resolve)."
    )
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    nr.italic = True

    out = r"C:\Users\Jace\Desktop\College Files\diko\dissertation\hardware_swot.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
